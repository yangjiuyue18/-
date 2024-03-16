import os
import requests
from bs4 import BeautifulSoup
from PIL import Image
from io import BytesIO
import PyPDF2
from pdf2image import convert_from_path
from docx import Document

def get_content_from_html(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')

    # 提取文本
    text = [p.text for p in soup.find_all('p')]

    # 提取图片
    images = []
    for img in soup.find_all('img'):
        if 'src' in img.attrs:
            img_url = img['src']
            img_response = requests.get(img_url)
            img_data = Image.open(BytesIO(img_response.content))
            images.append(img_data)

    return {'text': text, 'images': images}

def get_content_from_pdf(file_path):
    # 提取文本
    pdf_file_obj = open(file_path, 'rb')
    pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)
    text = []
    for page_num in range(pdf_reader.numPages):
        page_obj = pdf_reader.getPage(page_num)
        text.append(page_obj.extractText())
    pdf_file_obj.close()

    # 提取图片
    images = convert_from_path(file_path)

    return {'text': text, 'images': images}

def get_content_from_doc(file_path):
    # 提取文本
    doc = Document(file_path)
    text = [para.text for para in doc.paragraphs]

    # 提取图片
    images = []  # DOC文件的图片提取需要额外的处理

    return {'text': text, 'images': images}


def get_content_from_docx(file_path):
    # 提取文本
    doc = Document(file_path)
    text = [para.text for para in doc.paragraphs]

    # 提取图片
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_data = rel._target.blob
            images.append(image_data)

    return {'text': text, 'images': images}


def save_images(images, output_dir):
    for i, image_data in enumerate(images):
        # 使用PIL库打开图片
        image = Image.open(BytesIO(image_data))

        # 构造图片的文件路径
        image_file_path = os.path.join(output_dir, f'image_{i}.png')

        # 保存图片
        image.save(image_file_path)

def main(file_path, output_dir):
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    function_map = {
        '.html': get_content_from_html,
        '.pdf': get_content_from_pdf,
        '.doc': get_content_from_doc,
        '.docx': get_content_from_docx
    }

    if ext in function_map:
        content = function_map[ext](file_path)
        print('Text:', content['text'])
        print('Images:', len(content['images']))

        # 保存图片到本地
        save_images(content['images'], output_dir)
    else:
        print(f'Unsupported file type: {ext}')

if __name__ == "__main__":
    file_path = input('Enter the path to the file: ')
    output_dir = input('Enter the output directory: ')
    main(file_path, output_dir)